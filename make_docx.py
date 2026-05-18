#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""메리세이프 서비스 서류 통합본 Word 파일 생성"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─── 페이지 여백 설정 (A4) ───────────────────────────────────────────
for section in doc.sections:
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(1.8)
    section.right_margin  = Cm(1.8)

# ─── 기본 스타일 ───────────────────────────────────────────────────────
def set_font(run, size=9.5, bold=False, color=(74,74,74)):
    run.font.name = '맑은 고딕'
    run.font._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*color)

def add_para(doc_or_cell, text, size=9.5, bold=False, color=(74,74,74),
             align=WD_ALIGN_PARAGRAPH.LEFT, space_before=2, space_after=2,
             indent_left=0):
    p = doc_or_cell.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent_left:
        p.paragraph_format.left_indent = Cm(indent_left)
    p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return p

def add_section_title(doc, text, icon=''):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # 배경색 (짙은 회색)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1C1C1C')
    pPr.append(shd)
    run = p.add_run(f'  {icon}  {text}' if icon else f'  {text}')
    run.font.name = '맑은 고딕'
    run.font._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    return p

def add_doc_header(doc, title, subtitle='', is_first=False):
    if not is_first:
        # 페이지 나누기
        p = doc.add_paragraph()
        run = p.add_run()
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _OE
        br = _OE('w:br')
        br.set(_qn('w:type'), 'page')
        run._r.append(br)
    # 로고
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_before = Pt(0)
    p_logo.paragraph_format.space_after = Pt(2)
    r1 = p_logo.add_run('MARRYSAFE  ')
    r1.font.name = '맑은 고딕'
    r1.font._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    r1.font.size = Pt(16)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(0,0,0)
    r2 = p_logo.add_run('메리세이프')
    r2.font.name = '맑은 고딕'
    r2.font._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    r2.font.size = Pt(10)
    r2.font.bold = False
    r2.font.color.rgb = RGBColor(138,138,138)
    # 제목
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(2)
    p_title.paragraph_format.space_after = Pt(2)
    r = p_title.add_run(title)
    r.font.name = '맑은 고딕'
    r.font._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = RGBColor(28,28,28)
    # 부제목
    if subtitle:
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sub.paragraph_format.space_before = Pt(1)
        p_sub.paragraph_format.space_after = Pt(6)
        r = p_sub.add_run(subtitle)
        r.font.name = '맑은 고딕'
        r.font._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        r.font.size = Pt(8.5)
        r.font.bold = False
        r.font.color.rgb = RGBColor(100,100,100)
    # 구분선
    p_line = doc.add_paragraph()
    p_line.paragraph_format.space_before = Pt(2)
    p_line.paragraph_format.space_after = Pt(8)
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1C1C1C')
    pBdr.append(bottom)
    p_line._p.get_or_add_pPr().append(pBdr)

def set_table_style(table):
    table.style = 'Table Grid'
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after  = Pt(2)

def th_cell(cell, text, size=9, bold=True):
    cell.paragraphs[0].clear()
    pPr = cell.paragraphs[0]._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0EE')
    pPr.append(shd)
    run = cell.paragraphs[0].add_run(text)
    run.font.name = '맑은 고딕'
    run.font._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(28,28,28)

def td_cell(cell, text, size=9, bold=False, color=(74,74,74)):
    cell.paragraphs[0].clear()
    run = cell.paragraphs[0].add_run(text)
    run.font.name = '맑은 고딕'
    run.font._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*color)

def add_sign_row(doc, labels):
    """서명란 행 추가"""
    table = doc.add_table(rows=1, cols=len(labels))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_style(table)
    for i, label in enumerate(labels):
        cell = table.rows[0].cells[i]
        cell.paragraphs[0].clear()
        p1 = cell.paragraphs[0]
        r = p1.add_run(label)
        r.font.name = '맑은 고딕'
        r.font._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(138,138,138)
        p2 = cell.add_paragraph('\n\n')
        p2.paragraph_format.space_before = Pt(1)
        p2.paragraph_format.space_after = Pt(2)
    table.paragraph_format if hasattr(table,'paragraph_format') else None

def add_info_grid_table(doc, rows_data):
    """2열 기본 정보 그리드 (label + input line)"""
    t = doc.add_table(rows=len(rows_data), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_style(t)
    for i, (lbl, val, full) in enumerate(rows_data):
        if full:
            # 두 셀 병합
            cell = t.rows[i].cells[0].merge(t.rows[i].cells[1])
            _fill_info_cell(cell, lbl, val)
        else:
            _fill_info_cell(t.rows[i].cells[0], lbl, val)
            # 오른쪽 셀은 다음 항목이 있을 때 채워짐 (rows_data 쌍으로 구성됨)
    return t

def _fill_info_cell(cell, label, value):
    cell.paragraphs[0].clear()
    r_lbl = cell.paragraphs[0].add_run(label)
    r_lbl.font.name = '맑은 고딕'
    r_lbl.font._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    r_lbl.font.size = Pt(8)
    r_lbl.font.color.rgb = RGBColor(138,138,138)
    p2 = cell.add_paragraph(value)
    p2.paragraph_format.space_before = Pt(1)
    p2.paragraph_format.space_after = Pt(2)
    for run in p2.runs:
        run.font.name = '맑은 고딕'
        run.font._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(74,74,74)

def add_check_item(doc, text, sub='', indent=0.3):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(indent)
    r = p.add_run('□  ' + text)
    r.font.name = '맑은 고딕'
    r.font._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(28,28,28)
    if sub:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(1)
        p2.paragraph_format.left_indent  = Cm(indent + 0.6)
        r2 = p2.add_run(sub)
        r2.font.name = '맑은 고딕'
        r2.font._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        r2.font.size = Pt(8.5)
        r2.font.color.rgb = RGBColor(138,138,138)

def add_note_box(doc, text, style='tip'):
    """tip/info/warn 박스"""
    colors = {
        'tip':  (247,247,245),
        'info': (240,247,253),
        'warn': (253,240,240),
    }
    border_colors = {
        'tip':  '1C1C1C',
        'info': 'AACCE0',
        'warn': 'E0AAAA',
    }
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.3)
    r = p.add_run(text)
    r.font.name = '맑은 고딕'
    r.font._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    r.font.size = Pt(9)
    if style == 'warn':
        r.font.color.rgb = RGBColor(139,0,0)
    elif style == 'info':
        r.font.color.rgb = RGBColor(30,60,100)
    else:
        r.font.color.rgb = RGBColor(50,50,50)
    # 배경색
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    r_hex = '%02X%02X%02X' % colors[style]
    shd.set(qn('w:fill'), r_hex)
    pPr.append(shd)

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break()
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OE
    br = _OE('w:br')
    br.set(_qn('w:type'), 'page')
    run._r.append(br)


# ═══════════════════════════════════════════════════════════════
#  문서 1: 사전 체크리스트
# ═══════════════════════════════════════════════════════════════
add_doc_header(doc,
    '서비스 사전 체크리스트',
    '본 체크리스트는 서비스 이용계약서의 부속 서류로서 계약과 동일한 효력을 가집니다.\n예식일 3~5일 전까지 작성하여 제출해 주십시오.',
    is_first=True)

# A. 예식 기본 정보
add_section_title(doc, 'A. 예식 기본 정보', '📋')

t = doc.add_table(rows=7, cols=2)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_style(t)
rows_info = [
    ('고객 성명', ''),
    ('연락처', ''),
    ('예식 일시', '____년 ____월 ____일  ____시 ____분'),
    ('서비스 유형', '□ 4인2팀(양측)   □ 2인1팀(신랑측)   □ 2인1팀(신부측)'),
    ('예식장명', ''),
    ('홀 명', ''),
]
for i, (lbl, val) in enumerate(rows_info):
    _fill_info_cell(t.rows[i].cells[0], lbl, val)
    if i == 5:
        # 마지막 쌍
        _fill_info_cell(t.rows[i].cells[1], '', '')

# 7행: 예식장 주소 (병합)
cell_addr = t.rows[6].cells[0].merge(t.rows[6].cells[1])
_fill_info_cell(cell_addr, '예식장 주소', '')

# 더 행 추가
t2 = doc.add_table(rows=3, cols=2)
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_style(t2)
_fill_info_cell(t2.rows[0].cells[0], '예상 하객 수', '약 ________ 명')
_fill_info_cell(t2.rows[0].cells[1], '주차 정보', '')
_fill_info_cell(t2.rows[1].cells[0], '혼주 (신랑 측)', '성명: ________  ☎ 010-____-____')
_fill_info_cell(t2.rows[1].cells[1], '혼주 (신부 측)', '성명: ________  ☎ 010-____-____')
_fill_info_cell(t2.rows[2].cells[0], '메리세이프 담당자', '')
_fill_info_cell(t2.rows[2].cells[1], '담당자 연락처', '')

doc.add_paragraph()  # 간격

# B. 인수 대상자 정보
add_section_title(doc, 'B. 인수 대상자 정보', '👤')
add_para(doc, '서비스 종료 후 축의금 및 기록물을 인계받을 분의 정보를 기재합니다.\n인수 시 ① 신분증 실물 대조 + ② 사전 등록 전화번호 유선 확인 + ③ 인수인계 확인서 서명의 3단계 확인을 거칩니다.',
         size=9, color=(74,74,74))

tb = doc.add_table(rows=4, cols=5)
set_table_style(tb)
headers = ['구분','성명','성별','전화번호','신분증 종류']
for i, h in enumerate(headers):
    th_cell(tb.rows[0].cells[i], h)
rows_b = [
    ('1순위 인수자', '', '', '010 - ____ - ____', '□ 주민등록증  □ 운전면허증  □ 여권'),
    ('2순위 인수자\n(1순위 부재 시)', '', '', '010 - ____ - ____', '□ 주민등록증  □ 운전면허증  □ 여권'),
    ('비상 연락자\n(긴급 연락 전용)', '', '', '010 - ____ - ____', '승인 권한 없음, 긴급 연락 전용'),
]
for i, row_data in enumerate(rows_b):
    for j, val in enumerate(row_data):
        td_cell(tb.rows[i+1].cells[j], val)

doc.add_paragraph()

# C. 식권 배부 기준
add_section_title(doc, 'C. 식권 배부 기준', '🎫')

add_para(doc, 'C-1. 기본 배부 기준', size=9, bold=True)
tc1 = doc.add_table(rows=4, cols=2)
set_table_style(tc1)
th_cell(tc1.rows[0].cells[0], '항목')
th_cell(tc1.rows[0].cells[1], '고객 선택')
td_cell(tc1.rows[1].cells[0], '식권 배부 기준')
td_cell(tc1.rows[1].cells[1], '□ 축의금 납부자에 한해 배부   □ 모든 하객에게 배부   □ 기타: _______________')
td_cell(tc1.rows[2].cells[0], '1인당 식권 수량')
td_cell(tc1.rows[2].cells[1], '□ 1매   □ 2매   □ 봉투 수량대로   □ 기타: ______매')
td_cell(tc1.rows[3].cells[0], '잔여 식권 처리')
td_cell(tc1.rows[3].cells[1], '□ 인수 대상자에게 함께 전달   □ 예식장에 반환   □ 기타: _______________')

doc.add_paragraph()
add_para(doc, 'C-2. 축의금 없이 식권만 요청하는 하객 대응 (택 1)', size=9, bold=True)
tc2 = doc.add_table(rows=3, cols=2)
set_table_style(tc2)
th_cell(tc2.rows[0].cells[0], '상황')
th_cell(tc2.rows[0].cells[1], '고객 선택')
td_cell(tc2.rows[1].cells[0], '현금 없이 식권만 요청')
td_cell(tc2.rows[1].cells[1],
    '□ 식권 배부하지 않음\n'
    '□ 소란 방지를 위해 배부 (명단에 \'축의금 미납\' 기록)\n'
    '□ 혼주 확인 후 판단\n'
    '□ 기타: _______________________________________________')
td_cell(tc2.rows[2].cells[0], '봉투가 비어있는 경우')
td_cell(tc2.rows[2].cells[1],
    '□ 명단 기록, 식권 배부하지 않음\n'
    '□ 명단 기록, 식권 배부 (소란 방지)\n'
    '□ 혼주 즉시 유선 확인 후 판단\n'
    '□ 기타: _______________________________________________')
add_note_box(doc, '💡 메리세이프 권장사항: 혼잡한 웨딩 현장에서 분쟁이 생기면 모든 분이 불편해집니다. 특별한 사정이 없다면 「소란 방지를 위해 배부」를 권장합니다.', 'tip')

doc.add_paragraph()
add_para(doc, 'C-3. 계좌이체 주장 하객 응대 기준 (택 1)', size=9, bold=True)
tc3 = doc.add_table(rows=3, cols=2)
set_table_style(tc3)
th_cell(tc3.rows[0].cells[0], '항목')
th_cell(tc3.rows[0].cells[1], '고객 선택')
td_cell(tc3.rows[1].cells[0], '이체 내역 확인 여부')
td_cell(tc3.rows[1].cells[1],
    '□ 하객 이체 화면(캡처) 제시 요청 후 확인\n'
    '□ 이체 확인 없이 명단에 \'이체 주장\' 기록 후 배부\n'
    '□ 이체 확인 없이 식권 배부하지 않음\n'
    '□ 혼주 확인 후 판단')
td_cell(tc3.rows[2].cells[0], '이체 확인용 계좌\n(확인 선택 시 기재)')
td_cell(tc3.rows[2].cells[1], '예금주: ____________  은행: ____________  계좌번호: _______________________')
add_note_box(doc, '📌 메리세이프는 이체 사실의 진위를 최종 확인할 의무를 부담하지 않으며, 고객이 선택한 기준에 따라 응대합니다.', 'info')

doc.add_paragraph()

# D. 돌발상황 안내
add_section_title(doc, 'D. 돌발상황 안내 — 봉투 절취·교체 실제 사례', '⚠️')
add_para(doc, '웨딩 현장에서 실제로 발생한 사례입니다. 사전에 숙지하시면 피해 예방에 도움이 됩니다.', size=9, color=(100,100,100))

situations = [
    ('사례 ① 봉투 들고 사라지기',
     '접수대 혼잡 틈을 노려 하객으로 위장한 자가 테이블 위 봉투를 집어 들고 도주한 사례. 메리세이프는 봉투 수령 즉시 세이프 박스에 투입하여 원천 차단합니다.'),
    ('사례 ② 빈 봉투 교체 (바꿔치기)',
     '미리 준비한 동일한 외형의 빈 봉투(또는 소액 봉투)로 접수대 근처에서 바꿔치기 하는 사례. 피해액이 수십만 원에 달한 경우도 있습니다.'),
    ('사례 ③ "내가 맡아둔다"며 접근',
     '친척·지인으로 위장하여 "잠깐 보관하겠다"며 다른 하객의 봉투를 빼앗아 가는 수법. 접수 직원 외 제3자에게는 절대 봉투를 넘기지 않습니다.'),
    ('사례 ④ "혼주 가족이 보냈다"며 중간 인계 요청',
     '행사 도중 일부를 미리 달라고 요청하는 경우. 메리세이프는 사전 지정된 인수자 외에는 어떤 경우에도 중간 인계를 하지 않습니다.'),
]
for title, body in situations:
    td = doc.add_table(rows=2, cols=1)
    set_table_style(td)
    th_cell(td.rows[0].cells[0], title, size=9)
    td_cell(td.rows[1].cells[0], body, size=9)
    doc.add_paragraph()

add_note_box(doc,
    '🔒 세이프 박스 운영 원칙: 한 번 닫힌 세이프 박스는 웨딩이 완전히 종료될 때까지 절대 개봉하지 않습니다. '
    '혼주님 또는 가족분의 요청이 있더라도 예외 없이 적용됩니다.\n'
    '→ 행사 중 현금이 필요한 비상 상황을 대비해, 시작 전 별도 현금(예: 20~30만원)을 미리 준비해 두시길 권장합니다.',
    'warn')

doc.add_paragraph()

# E. 기타 특별 요청 사항
add_section_title(doc, 'E. 기타 특별 요청 사항', '📝')
te = doc.add_table(rows=4, cols=2)
set_table_style(te)
th_cell(te.rows[0].cells[0], 'No', size=9)
th_cell(te.rows[0].cells[1], '특별 요청 내용 (서비스 범위 내 수용 가능한 사항에 한함)', size=9)
for i in range(1, 4):
    td_cell(te.rows[i].cells[0], str(i))
    td_cell(te.rows[i].cells[1], '\n')

doc.add_paragraph()

# F. 확인 및 서명
add_section_title(doc, 'F. 확인 및 서명', '✍️')
add_para(doc, '위 체크리스트 내용을 확인하였으며, 기재되지 않은 상황 발생 시 회사의 표준 운영 기준에 따릅니다.', size=9, color=(100,100,100))

tf = doc.add_table(rows=1, cols=2)
set_table_style(tf)
_fill_info_cell(tf.rows[0].cells[0], '고객 (신랑 측)', '성명: _________________   서명(인):\n\n\n')
_fill_info_cell(tf.rows[0].cells[1], '고객 (신부 측)', '성명: _________________   서명(인):\n\n\n')
doc.add_paragraph()
tf2 = doc.add_table(rows=1, cols=2)
set_table_style(tf2)
_fill_info_cell(tf2.rows[0].cells[0], '작성일', '____년 ____월 ____일\n')
_fill_info_cell(tf2.rows[0].cells[1], '메리세이프 담당자', '성명: _________________   서명(인):\n\n\n')


# ═══════════════════════════════════════════════════════════════
#  문서 2: 개인정보 수집·이용 동의서
# ═══════════════════════════════════════════════════════════════
add_doc_header(doc, '개인정보 수집·이용 동의서')

add_section_title(doc, '1. 수집하는 개인정보 항목', '①')
t2_1 = doc.add_table(rows=4, cols=3)
set_table_style(t2_1)
th_cell(t2_1.rows[0].cells[0], '구분')
th_cell(t2_1.rows[0].cells[1], '수집 대상')
th_cell(t2_1.rows[0].cells[2], '수집 항목')
rows_2_1 = [
    ('필수','고객(신랑·신부/혼주)','성명, 연락처, 예식 일시·장소, 이메일'),
    ('필수','인수 대상자·비상 연락자','성명, 성별, 연락처, 신분증 정보(확인용, 사본 미보관)'),
    ('필수','하객','성명, 영상 기록(접수 과정 촬영분)'),
]
for i, (a,b,c) in enumerate(rows_2_1):
    td_cell(t2_1.rows[i+1].cells[0], a)
    td_cell(t2_1.rows[i+1].cells[1], b)
    td_cell(t2_1.rows[i+1].cells[2], c)
add_para(doc, '※ 봉투 내 금액은 확인·기록하지 않으므로 수집하지 않습니다.', size=8.5, color=(138,138,138))

doc.add_paragraph()
add_section_title(doc, '2. 수집·이용 목적', '②')
t2_2 = doc.add_table(rows=5, cols=2)
set_table_style(t2_2)
th_cell(t2_2.rows[0].cells[0], '목적')
th_cell(t2_2.rows[0].cells[1], '상세 내용')
rows_2_2 = [
    ('서비스 이행','축의금 접수, 명단 기록, 보관, 식권 배부 보조 등 계약상 서비스 수행'),
    ('인수인계','인수 대상자 본인 확인 및 축의금·기록물의 안전한 인계'),
    ('증빙 보존','서비스 수행 과정 투명성 확보 및 사후 분쟁 발생 시 증빙 자료 활용'),
    ('고객 연락','사전 확인, 현장 긴급 연락, 사후 문의 대응'),
]
for i, (a,b) in enumerate(rows_2_2):
    td_cell(t2_2.rows[i+1].cells[0], a)
    td_cell(t2_2.rows[i+1].cells[1], b)

doc.add_paragraph()
add_section_title(doc, '3. 개인정보 보유 및 파기', '③')
t2_3 = doc.add_table(rows=5, cols=3)
set_table_style(t2_3)
th_cell(t2_3.rows[0].cells[0], '정보 유형')
th_cell(t2_3.rows[0].cells[1], '보유 기간')
th_cell(t2_3.rows[0].cells[2], '파기 방법')
rows_2_3 = [
    ('고객·인수자 정보','예식일로부터 90일','전자 파일: 복구 불가 방법 영구 삭제 / 종이: 분쇄'),
    ('하객 명단','예식일로부터 30일','전자 파일: 복구 불가 방법 영구 삭제'),
    ('영상 기록물','예식일로부터 30일 (분쟁 시 종결까지)','전자 파일: 복구 불가 방법 영구 삭제'),
    ('계약 관련 서류','계약 종료 후 5년 (상법 제64조)','보유 기간 경과 후 분쇄 또는 영구 삭제'),
]
for i, (a,b,c) in enumerate(rows_2_3):
    td_cell(t2_3.rows[i+1].cells[0], a)
    td_cell(t2_3.rows[i+1].cells[1], b)
    td_cell(t2_3.rows[i+1].cells[2], c)

doc.add_paragraph()
add_section_title(doc, '4. 제3자 제공 및 동의 거부 권리', '④')
add_para(doc, '① 회사는 원칙적으로 수집한 개인정보를 제3자에게 제공하지 않습니다. 단, 고객 사전 동의 또는 법령에 의한 경우는 예외입니다.', size=9.5)
add_para(doc, '② 고객은 개인정보 수집·이용에 대한 동의를 거부할 권리가 있습니다. 단, 필수 항목 동의 거부 시 서비스 이용이 제한됩니다.', size=9.5)

doc.add_paragraph()
# 동의 체크
p_agree = doc.add_paragraph()
p_agree.paragraph_format.space_before = Pt(4)
p_agree.paragraph_format.space_after = Pt(4)
p_agree.paragraph_format.left_indent = Cm(0.3)
r = p_agree.add_run('□  개인정보 수집·이용에 동의합니다.          □  동의하지 않습니다.')
r.font.name = '맑은 고딕'
r.font._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
r.font.size = Pt(10)
r.font.bold = True
r.font.color.rgb = RGBColor(28,28,28)

doc.add_paragraph()
ts2 = doc.add_table(rows=1, cols=3)
set_table_style(ts2)
_fill_info_cell(ts2.rows[0].cells[0], '고객 성명', '\n\n\n')
_fill_info_cell(ts2.rows[0].cells[1], '서명(인)', '\n\n\n')
_fill_info_cell(ts2.rows[0].cells[2], '작성일', '____년 ____월 ____일\n')


# ═══════════════════════════════════════════════════════════════
#  문서 3: 촬영(녹화) 안내 및 동의서
# ═══════════════════════════════════════════════════════════════
add_doc_header(doc, '촬영(녹화) 안내 및 동의서',
    '중요: 메리세이프 서비스는 투명한 서비스 제공을 위해 현장에 카메라를 설치합니다.\n동의하지 않으시는 경우 서비스 이용이 불가합니다.')

add_section_title(doc, '1. 촬영 목적', '①')
t3_1 = doc.add_table(rows=4, cols=2)
set_table_style(t3_1)
th_cell(t3_1.rows[0].cells[0], '목적')
th_cell(t3_1.rows[0].cells[1], '상세 내용')
rows_3_1 = [
    ('투명성 확보','축의금 접수·보관·인수인계 전 과정을 영상으로 기록하여 서비스 수행의 투명성 보장'),
    ('증빙 자료','봉투 수령·보관·전달 과정을 객관적으로 증명할 수 있는 증거 확보'),
    ('분쟁 예방','봉투 누락, 인수인계 관련 분쟁 발생 시 사실 확인을 위한 증빙으로 활용'),
]
for i, (a,b) in enumerate(rows_3_1):
    td_cell(t3_1.rows[i+1].cells[0], a)
    td_cell(t3_1.rows[i+1].cells[1], b)

doc.add_paragraph()
add_section_title(doc, '2. 촬영 범위 및 방식', '②')
clauses = [
    ('①', '촬영 장소: 축의금 접수 테이블 및 그 주변 (잠금 박스 보관 영역 포함)'),
    ('②', '촬영 대상: 봉투 수령·처리·보관·인수인계 과정 (하객의 모습이 부수적으로 포함될 수 있음)'),
    ('③', '촬영 장비: 카메라 2대 (클로즈업·미디엄 각 1대)'),
    ('④', '촬영 시간: 예식 시작 1시간 전부터 인수 대상자에게 축의금이 전달되는 시점까지'),
    ('⑤', '영상 제공: 파일 형태의 제공은 불가하며, 유튜브(YouTube) 비공개 링크로만 제공 (열람 기간 제한 없음)'),
    ('⑥', '고지: 접수 테이블 주변에 "보안 촬영 중" 안내문을 게시하여 하객에게 촬영 사실을 고지합니다.'),
]
for num, text in clauses:
    add_para(doc, f'{num}  {text}', size=9.5)

doc.add_paragraph()
add_section_title(doc, '3. 영상 기록물 보관·파기 및 이용 제한', '③')
t3_3 = doc.add_table(rows=5, cols=2)
set_table_style(t3_3)
th_cell(t3_3.rows[0].cells[0], '항목')
th_cell(t3_3.rows[0].cells[1], '내용')
rows_3_3 = [
    ('보관 기간','예식일로부터 30일 (분쟁 발생 시 해당 분쟁 종결 시까지 연장)'),
    ('보관 방법','유튜브 비공개 저장 + 로컬 SSD 이중 보관, 접근 권한 제한'),
    ('파기 방법','보관 기간 경과 후 복구 불가능한 방법으로 영구 삭제'),
    ('이용 제한','촬영 목적 외 사용 금지 / 제3자 제공 금지 / 마케팅·홍보 목적 사용 금지'),
]
for i, (a,b) in enumerate(rows_3_3):
    td_cell(t3_3.rows[i+1].cells[0], a)
    td_cell(t3_3.rows[i+1].cells[1], b)

doc.add_paragraph()
p_agree3 = doc.add_paragraph()
p_agree3.paragraph_format.space_before = Pt(4)
p_agree3.paragraph_format.space_after = Pt(4)
p_agree3.paragraph_format.left_indent = Cm(0.3)
r3 = p_agree3.add_run('□  촬영(녹화)에 동의합니다. (동의하지 않을 경우 서비스 이용이 불가합니다.)')
r3.font.name = '맑은 고딕'
r3.font._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
r3.font.size = Pt(10)
r3.font.bold = True
r3.font.color.rgb = RGBColor(28,28,28)

doc.add_paragraph()
ts3 = doc.add_table(rows=1, cols=3)
set_table_style(ts3)
_fill_info_cell(ts3.rows[0].cells[0], '고객 성명', '\n\n\n')
_fill_info_cell(ts3.rows[0].cells[1], '서명(인)', '\n\n\n')
_fill_info_cell(ts3.rows[0].cells[2], '작성일', '____년 ____월 ____일\n')


# ═══════════════════════════════════════════════════════════════
#  문서 4: 인수인계 확인서
# ═══════════════════════════════════════════════════════════════
add_doc_header(doc, '인수인계 확인서',
    '서비스 종료 시 축의금 및 기록물의 인계가 이상 없이 완료되었음을 상호 확인하기 위하여 작성합니다.\n예식 당일 서비스 종료 후 작성합니다.')

# A. 서비스 기본 정보
add_section_title(doc, 'A. 서비스 기본 정보', 'A')
ta = doc.add_table(rows=4, cols=2)
set_table_style(ta)
_fill_info_cell(ta.rows[0].cells[0], '고객 성명', '')
_fill_info_cell(ta.rows[0].cells[1], '계약번호', '')
_fill_info_cell(ta.rows[1].cells[0], '예식 일시', '____년 ____월 ____일  ____시 ____분')
_fill_info_cell(ta.rows[1].cells[1], '예식장 / 홀', '')
_fill_info_cell(ta.rows[2].cells[0], '서비스 유형', '□ 4인2팀(양측)   □ 2인1팀(신랑측)   □ 2인1팀(신부측)')
_fill_info_cell(ta.rows[2].cells[1], '운영 인력 수', '________ 명')
_fill_info_cell(ta.rows[3].cells[0], '서비스 시작', '____시 ____분')
_fill_info_cell(ta.rows[3].cells[1], '서비스 종료', '____시 ____분')

doc.add_paragraph()

# B. 인수 대상자 본인 확인
add_section_title(doc, 'B. 인수 대상자 본인 확인', 'B')
tb4 = doc.add_table(rows=2, cols=4)
set_table_style(tb4)
th_cell(tb4.rows[0].cells[0], '성명')
th_cell(tb4.rows[0].cells[1], '신분증 확인')
th_cell(tb4.rows[0].cells[2], '유선 확인')
th_cell(tb4.rows[0].cells[3], '인수자 순위')
td_cell(tb4.rows[1].cells[0], '\n')
td_cell(tb4.rows[1].cells[1], '□ 확인 완료')
td_cell(tb4.rows[1].cells[2], '□ 확인 완료')
td_cell(tb4.rows[1].cells[3], '□ 1순위   □ 2순위   □ 기타')

doc.add_paragraph()

# C. 인계 물품 내역
add_section_title(doc, 'C. 인계 물품 내역', 'C')
tc4 = doc.add_table(rows=5, cols=4)
set_table_style(tc4)
th_cell(tc4.rows[0].cells[0], 'No')
th_cell(tc4.rows[0].cells[1], '인계 항목')
th_cell(tc4.rows[0].cells[2], '수량/내용')
th_cell(tc4.rows[0].cells[3], '확인')
rows_c = [
    ('1','축의금 봉투 (잠금 박스 포함)','총 ________ 개','□ 인수자 확인'),
    ('2','하객 명단 기록물','□ 전산파일  □ 출력물','□ 인수자 확인'),
    ('3','잔여 식권 (해당 시)','총 ________ 매','□ 인수자 확인'),
    ('4','기타','','□ 인수자 확인'),
]
for i, (a,b,c,d) in enumerate(rows_c):
    td_cell(tc4.rows[i+1].cells[0], a)
    td_cell(tc4.rows[i+1].cells[1], b)
    td_cell(tc4.rows[i+1].cells[2], c)
    td_cell(tc4.rows[i+1].cells[3], d)
add_para(doc, '※ 금액 기록은 서비스 범위에 포함되지 않으므로 정산 내역 인계는 없습니다.', size=8.5, color=(138,138,138))

doc.add_paragraph()

# D. 서비스 이행 확인
add_section_title(doc, 'D. 서비스 이행 확인', 'D')
td4 = doc.add_table(rows=6, cols=3)
set_table_style(td4)
th_cell(td4.rows[0].cells[0], '확인 항목')
th_cell(td4.rows[0].cells[1], '결과')
th_cell(td4.rows[0].cells[2], '비고')
rows_d = [
    ('축의금 접수·보관이 계약 내용에 따라 수행되었는가','□ 이상 없음   □ 이상 있음',''),
    ('영상 녹화가 정상적으로 진행되었는가','□ 이상 없음   □ 이상 있음',''),
    ('명단 기록이 정상적으로 작성되었는가','□ 이상 없음   □ 이상 있음',''),
    ('인수인계 시 봉투 수량과 기록이 일치하는가','□ 이상 없음   □ 이상 있음',''),
    ('서비스 수행 중 특이사항이 있었는가','□ 없음   □ 있음 (E항 기재)',''),
]
for i, (a,b,c) in enumerate(rows_d):
    td_cell(td4.rows[i+1].cells[0], a)
    td_cell(td4.rows[i+1].cells[1], b)
    td_cell(td4.rows[i+1].cells[2], c)

doc.add_paragraph()

# E. 특이사항
add_section_title(doc, 'E. 특이사항', 'E')
te4 = doc.add_table(rows=1, cols=1)
set_table_style(te4)
td_cell(te4.rows[0].cells[0], '\n\n\n\n')

doc.add_paragraph()

# F. 최종 확인 및 서명
add_section_title(doc, 'F. 최종 확인 및 서명', 'F')
add_para(doc,
    '위 내용을 확인하며, 축의금 및 기록물의 인계가 이상 없이 완료되었음을 상호 확인합니다. '
    '본 확인서 작성 이후 축의금 및 기록물에 대한 관리 책임은 인수자에게 이전됩니다.',
    size=9, color=(100,100,100))
tf4 = doc.add_table(rows=1, cols=2)
set_table_style(tf4)
_fill_info_cell(tf4.rows[0].cells[0], '회사 (인계자)', '성명: _____________   서명(인):\n\n\n')
_fill_info_cell(tf4.rows[0].cells[1], '인수자', '성명: _____________   서명(인):\n\n\n')

doc.add_paragraph()
add_para(doc, '인수인계 일시: ____년 ____월 ____일  ____시 ____분    /    장소: _______________________________',
         size=9, color=(74,74,74))

doc.add_paragraph()
# 푸터
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_footer.paragraph_format.space_before = Pt(10)
r = p_footer.add_run('© 2025 MarrySafe (메리세이프)  ·  marrysafe.official@gmail.com  ·  marrysafe.co.kr\n본 서류는 서비스 계약과 함께 보관됩니다.')
r.font.name = '맑은 고딕'
r.font._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
r.font.size = Pt(8)
r.font.color.rgb = RGBColor(138,138,138)

# ─── 저장 ────────────────────────────────────────────────────────────
output_path = '/Users/younglee/Desktop/메리세이프 홈페이지/메리세이프_서비스서류_통합본.docx'
doc.save(output_path)
print(f'✅ 저장 완료: {output_path}')
